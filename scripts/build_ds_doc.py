"""
Build documentos/ds_mlops_portfolio.docx — extra RAG context that emphasises
Domingo's Data Science / Data Pipelines / MLOps work, the tooling used in
this portfolio project, and the broader DS toolkit. Run once after editing.

This is a separate generator (not modifying scripts/index_documents.py) so
the existing pipeline keeps doing exactly what it did before; we only add
a new source document for it to ingest.

Usage:
    python scripts/build_ds_doc.py
"""

from pathlib import Path
from docx import Document

OUT = Path(__file__).resolve().parents[1] / "documentos" / "ds_mlops_portfolio.docx"

SECTIONS = [
    ("Perfil profesional — Data Science / MLOps",
     [
         "Domingo Berbel es Data Scientist con foco en pipelines de datos, modelos predictivos en producción y MLOps. Trabaja end-to-end: ingesta, transformación, entrenamiento, despliegue, observabilidad y mantenimiento.",
         "Combina rol técnico (construcción de pipelines, modelos, infraestructura cloud) con rol analítico (visualización, comunicación de resultados a negocio).",
         "Áreas de especialidad: data engineering en Spark/PySpark, modelos de clasificación y series temporales, sistemas de recomendación, retrieval-augmented generation (RAG) sobre LLMs, y automatización de pipelines en GitHub Actions y Azure.",
     ]),

    ("Proyecto portfolio — domingoberbel.com",
     [
         "Web personal y de portfolio desplegada en Azure (Static Web Apps + Container Apps). Demuestra el ciclo completo de Data Science: ingesta de datos en tiempo real, ETL, entrenamiento de modelos, exposición de resultados como visualizaciones interactivas y un chatbot RAG sobre el CV.",
         "Stack: Python (FastAPI, scikit-learn, pandas, numpy, PySpark), JavaScript (React, Vite, Recharts, Tailwind), infraestructura (Docker, GitHub Actions, Azure Container Apps, Azure Static Web Apps, GitHub Container Registry, Hostinger DNS).",
         "Modelos integrados en la web: predicción de La Liga (Poisson + Monte Carlo + XGBoost), credit scoring (Logistic Regression + Gradient Boosting calibrado), recomendador de productos content-based con TF-IDF y MMR, y un chatbot conversacional RAG sobre los documentos del CV.",
     ]),

    ("Pipeline de datos — La Liga (ETL en producción)",
     [
         "Pipeline ETL programado cada 30 minutos vía cron de GitHub Actions (schedule: '*/30 * * * *' en .github/workflows/update-laliga.yml). Es producción real, no maqueta: corre desde 2025 sin intervención manual.",
         "Extract: llamada a la API pública SportsRC (https://api.sportsrc.org/) que devuelve clasificación, partidos y goleadores de La Liga en JSON.",
         "Transform: PySpark con DataFrames y schemas explícitos (StructType). Calcula ratios por partido (ppg, winRate, drawRate, lossRate, gfPerGame, gaPerGame, gdPerGame) y clasifica equipos en zonas (Champions, Europa, Mid, Relegation) con expresiones F.when() encadenadas.",
         "Load: persiste el resultado en frontend/public/data/la_liga_data.json y lo commitea de vuelta al repositorio. El frontend lo consume directamente como archivo estático servido por Azure Static Web Apps. Latencia de visualización para el usuario: cero (CDN-cached).",
         "Modelo predictivo (scripts/predict_laliga.py): feature engineering, simulación Poisson de cada partido restante, Monte Carlo con 1.000 temporadas simuladas, y XGBoost multi:softprob para clasificar zonas. Output: la_liga_predictions.json con probabilidades de campeón, máximo goleador, descenso, etc.",
     ]),

    ("Modelo de credit scoring — banca",
     [
         "Pipeline de clasificación binaria (default / no-default) implementado en scripts/credit_scoring_pipeline.py. Genera dataset sintético calibrado (5.000 solicitantes, ~17% tasa de default), preprocesa con ColumnTransformer (StandardScaler + OneHotEncoder), y entrena dos modelos en paralelo: Logistic Regression como modelo interpretable y Gradient Boosting Classifier como modelo de mayor capacidad.",
         "Validación rigurosa: split estratificado 75/25, validación cruzada estratificada 5-fold, métricas en held-out test set. Reporta AUC-ROC, GINI, KS-statistic, Average Precision, F1, Precision, Recall y Brier score. Calibración del Gradient Boosting vía Platt scaling (sigmoid) para que las probabilidades reflejen la realidad.",
         "Scorecard PDO 300-850 estilo FICO: factor = PDO / ln(2), offset = base_score − factor · ln(base_odds). Bandas Poor/Fair/Good/Very Good/Exceptional. Permutation feature importance (model-agnostic) sobre el test set para identificar las features más predictivas.",
         "Inferencia interactiva en navegador: los coeficientes y parámetros del scaler de la regresión logística se exportan a JSON, y el navegador ejecuta sigmoid(β·x) directamente en JavaScript. Latencia cero de red, sin coste de cómputo en backend, sliders interactivos en el frontend que recalculan score y probabilidad de default en vivo.",
     ]),

    ("Sistema de recomendación — content-based",
     [
         "Recomendador content-based implementado en scripts/recommendation_engine.py. Catálogo de 32 productos en 4 categorías (electronics, books, sports, home). Cada producto se vectoriza combinando TF-IDF sobre la descripción (1-2 grams, max 120 features, stop-words EN), one-hot de categoría, y MinMax-scaler de precio y rating. Vector resultante de 126 dimensiones.",
         "Perfil del usuario = media de los vectores de los productos seleccionados. Score = cosine similarity entre perfil y cada producto. Re-ranking MMR (Maximal Marginal Relevance) con λ=0.7 para equilibrar relevancia y diversidad y evitar listas redundantes.",
         "Métricas: intra-list diversity (1 − media de similitud entre pares recomendados) y catalog coverage (% del catálogo cubierto). Evaluado sobre 4 personas sintéticas (Tech Enthusiast, Data Scientist, Fitness Fan, Home Professional).",
         "Inferencia client-side: la matriz de features completa, el vocabulario TF-IDF y los metadatos del catálogo se exportan a JSON. El frontend implementa cosine similarity, mean vector y MMR en JavaScript puro. Permite explicabilidad: al seleccionar una recomendación, el frontend descompone la similaridad en contribuciones (TF-IDF términos, categoría, precio/rating) y muestra los términos del vocabulario más responsables de la elección.",
     ]),

    ("RAG — Retrieval-Augmented Generation sobre el CV",
     [
         "Chatbot conversacional con RAG implementado en backend/app/services/rag_service.py. Indexa los documentos del CV de Domingo (cv_rag.docx, letter_rag.docx, more_information.docx) en chunks de 650 caracteres con 100 de overlap, y genera embeddings de 3.072 dimensiones con Google Gemini (gemini-embedding-001).",
         "Recuperación híbrida: combina similitud por coseno sobre los embeddings (semántica) con BM25 (palabras clave) y los fusiona vía Reciprocal Rank Fusion (RRF). Evita el problema de que la búsqueda semántica pura ignore términos exactos importantes (nombres propios, siglas).",
         "Generación: una vez recuperados los chunks relevantes, se construye un prompt con system prompt + contexto + pregunta y se envía a OpenAI/Azure OpenAI. Respuesta streamed por Server-Sent Events (SSE) al frontend para que aparezca token a token.",
         "Guardrails: rate limiting por IP, captcha opcional (Cloudflare Turnstile), token budget máximo por sesión, y filtrado de preguntas fuera de tema (saludos, contacto, inapropiadas) con respuestas rápidas pre-construidas. Logging de todas las preguntas para análisis posterior (analytics_store.py + SQLite).",
     ]),

    ("Stack de Data Science y MLOps utilizado",
     [
         "Lenguajes y entornos: Python 3.10+, JavaScript (ES2022), SQL, Bash. IDEs y notebooks: VS Code, Jupyter.",
         "Procesamiento de datos: pandas, numpy, PySpark (DataFrames con schema explícito), Apache Spark, SQL.",
         "Machine Learning clásico: scikit-learn (LogisticRegression, GradientBoostingClassifier, RandomForest, ColumnTransformer, StandardScaler, OneHotEncoder, MinMaxScaler, TfidfVectorizer, StratifiedKFold, train_test_split, permutation_importance, CalibratedClassifierCV), XGBoost (multi:softprob), LightGBM.",
         "Métricas: AUC-ROC, GINI, KS-statistic, F1, precision, recall, Average Precision, Brier score, log-loss, MAE/MAPE, intra-list diversity, catalog coverage.",
         "NLP y embeddings: TF-IDF (n-grams, stop-words, sub-linear TF), Sentence-Transformers, Google Gemini embeddings (gemini-embedding-001, 3.072 dims), OpenAI embeddings, BM25 (rank-bm25), Reciprocal Rank Fusion (RRF).",
         "LLMs y agentes: OpenAI GPT-4 / GPT-4o-mini, Azure OpenAI, Google Gemini, prompt engineering, function calling, RAG, Server-Sent Events para streaming.",
         "Visualización: Recharts (frontend), matplotlib, seaborn, Power BI (avanzado), DAX, dashboards interactivos.",
         "Big data y data warehousing: Snowflake (certificación SnowPro Associate: Platform), Databricks (Fundamentals + GenAI Fundamentals), Delta Lake, Lakehouse architecture.",
         "MLOps e infraestructura: GitHub Actions (CI/CD, cron jobs, deploy automatizado a Azure SWA), Docker, GitHub Container Registry (GHCR), Azure Container Apps, Azure Static Web Apps, monitoring de cold-start, presupuestos y alertas en Azure.",
         "Backend y APIs: FastAPI, Uvicorn, Pydantic, Starlette, httpx, REST, SSE, OpenAPI/Swagger.",
         "Frontend y dataviz: React, Vite, Tailwind CSS, Recharts, Material Symbols.",
         "Bases de datos y caching: SQLite (analytics + logs), JSON file caching, Redis-style rate limiting in-memory.",
         "Validación y monitoreo: 5-fold StratifiedKFold, held-out test sets, Platt sigmoid calibration, Brier score para calibración, permutation importance, ROC y precision-recall curves.",
     ]),

    ("Patrones MLOps demostrados en este portfolio",
     [
         "Pipelines reproducibles: cada modelo (predict_laliga, credit_scoring_pipeline, recommendation_engine) parte de seeds fijos y produce outputs deterministas, controlables en versión.",
         "Separación entre entrenamiento e inferencia: los modelos se entrenan offline (en local o GitHub Actions) y se sirven como JSON estático. Para los modelos lineales se exportan los coeficientes y se reimplementa la inferencia en el navegador, eliminando latencia de red y coste de servidor.",
         "Cron jobs para refresco de datos: GitHub Actions con schedule que ejecuta el ETL cada 30 minutos, commitea los nuevos datos al repo, dispara el build del frontend y el deploy a Azure SWA. Sin intervención manual.",
         "Observabilidad y guardrails: rate limiting, token budgets, alertas de presupuesto en Azure, logging persistente de todas las consultas al chatbot, y health checks que despiertan el container en cold-start.",
         "Calibración de probabilidades: el credit scoring usa CalibratedClassifierCV con sigmoid (Platt scaling) para que las probabilidades sean fiables, no solo el ranking.",
         "Caching inteligente: embeddings cache (embeddings_cache.json) con metadata para regeneración automática solo si cambian los documentos o el modelo de embedding.",
     ]),

    ("Certificaciones relevantes (Data Science / Data Engineering)",
     [
         "Snowflake — SnowPro Associate: Platform (Oct 2025, expira Oct 2027). Cubre Snowflake, data warehousing, cloud, SQL avanzado.",
         "Databricks — Databricks Fundamentals Accreditation (Nov 2025). Cubre Databricks, Spark, Lakehouse, Delta Lake.",
         "Databricks — Generative AI Fundamentals (May 2025, expira May 2027). Cubre GenAI, LLMs, RAG, prompt engineering.",
         "Kaggle — Advanced SQL (Aug 2025). Window functions, CTEs, analytics queries.",
         "IBM — Data Analysis Using Python (2025). Pandas, NumPy, scikit-learn.",
         "LinkedIn Learning — Power BI Avanzado (May 2024). DAX, data modeling, dashboards.",
     ]),
]


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Domingo Berbel — Data Science, Pipelines y MLOps (extra RAG context)", level=1)
    doc.add_paragraph(
        "Este documento aporta contexto adicional al chatbot RAG sobre el trabajo "
        "de Data Science y MLOps de Domingo, las herramientas que usa y los modelos "
        "construidos en este propio portfolio (domingoberbel.com)."
    )

    for title, paragraphs in SECTIONS:
        doc.add_heading(title, level=2)
        for p in paragraphs:
            doc.add_paragraph(p)

    doc.save(OUT)
    print(f"Wrote {OUT} ({OUT.stat().st_size} bytes, {len(SECTIONS)} sections)")


if __name__ == "__main__":
    main()
